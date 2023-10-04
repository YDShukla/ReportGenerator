# Import necessary libraries

import os

from flask import Flask, render_template, session, redirect, request, send_file, render_template_string, url_for

import pdfkit

import io

from docx import Document

import tempfile

import openai

import xlsxwriter

import openpyxl

import dotenv

import re

import pandas as pd

import numpy as np

import json

from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime

from sqlalchemy.ext.declarative import declarative_base

from sqlalchemy.orm import sessionmaker

from datetime import datetime

dotenv.load_dotenv()

 

# Create a Flask web application

app = Flask(__name__, static_url_path='/static')

 

# Set the secret key for session management

app.secret_key = "your_secret_key"  # Replace with your own secret key

 

# Simulated user data (replace with a real authentication system)

users = {'user1': 'password1', 'user2': 'password2'}

admin_user = {'AD': 'BC'}

 

# SQLAlchemy setup

Base = declarative_base()

db_url = 'sqlite:///weekreport.db'  # SQLite database URL

engine = create_engine(db_url, echo=True)

 

# Define the data model for the SQLite database

class WeekReport(Base):

    __tablename__ = 'weekreport'

 

    id = Column(Integer, primary_key=True)

    date_column = Column(DateTime)

    username = Column(String(255))

    input_ = Column(Text)

    output_ = Column(Text)

    business_update = Column(Text)

    service = Column(String(255))

    portfolio = Column(String(255))

    teammates = Column(String(255))

    progress = Column(String(255))

 

# Create the database tables

Base.metadata.create_all(engine)

 

# Create a session to interact with the database

Session = sessionmaker(bind=engine)

db_session = Session()

 

# Initialize the OpenAI API client

openai.api_key = os.getenv('OPEN_AI_API')

 

# Define a route for the homepage

@app.route('/')

def login():

    return render_template('login.html')

 

# Define a route for user authentication

@app.route('/login', methods=['POST'])

def authenticate():

    username = request.form['username']

    password = request.form['password']

 

    if username in users and users[username] == password:

        session['username'] = username

        session['submission'] = ''

        session['gpt_response'] = ''

        return render_template('form.html', submission='', username=username)

    elif username in admin_user and admin_user[username] == password:

        session['username'] = username

        session['submission'] = ''

        session['gpt_response'] = ''

        return render_template('admin_landing_page.html', submission='', username=username)

    else:

        return "Authentication failed. <a href='/'>Try again</a>"

 

# Define a route for the form submission page

@app.route('/process_form', methods=['POST'])

def process_form():

    # Extract form data (input fields from your HTML form)

    input_data = request.form['work']

    portfolio = request.form['project']

    service = request.form['services']

    selected_date = request.form['selected_date']

    progress = request.form['progress']

    team = request.form['team']

 

    # Use the OpenAI API to generate a response

    prompts = [

        {

            'role': 'system',

            'content': """

            You are a professional business report generator. Your task is to create a detailed business report in the following format, which includes sections for input, output, and a business update. Maintain a high level of professionalism in the language and presentation of the report.

 

            Please adhere to the specific format provided below:

 

            INPUT:

            Generate a 3-4 word long name for a response that focuses on the work done this week. Be concise.

 

            OUTPUT:

            Generate a concise one-line response that focuses on the outcome of the work done this week. Highlight aspects such as efficiency gains, reduced efforts, time savings, or other relevant results.

 

            BUSINESS UPDATE:

            Generate a succinct one-line statement that focuses on the updates related to the business from the generated output. Discuss how efficiency is improved, efforts are reduced, or any other pertinent updates that align with the organization's goals and objectives.

            """

        },

        {

            'role': 'user',

            'content': f"{input_data}"

        }

    ]

 

    response = openai.ChatCompletion.create(

        model="gpt-3.5-turbo",

        messages=prompts,

        temperature=0.0

    )

 

    gpt_response = response.choices[0].message.content

 

    # Store session data

    session['submission'] = gpt_response

    session['gpt_response'] = gpt_response

    session['portfolio'] = portfolio

    session['service'] = service

    session['selected_date'] = selected_date

    session['progress'] = progress

    session['team'] = team

 

    return redirect(url_for('submission_output_editable'))

 

# Define a route for displaying the submission output (editable)

@app.route('/submission_output_editable')

def submission_output_editable():

    text = session.get('gpt_response', '')

 

    input_pattern = r'INPUT:(.*?)OUTPUT:'

    output_pattern = r'OUTPUT:(.*?)BUSINESS UPDATE:'

    business_update_pattern = r'BUSINESS UPDATE:(.*)'

 

    # Use re.DOTALL to match across multiple lines

    input_match = re.search(input_pattern, text, re.DOTALL)

    output_match = re.search(output_pattern, text, re.DOTALL)

    business_update_match = re.search(business_update_pattern, text, re.DOTALL)

 

    # Extract the matched content

    input_section = input_match.group(1).strip() if input_match else ""

    output_section = output_match.group(1).strip() if output_match else ""

    business_update_section = business_update_match.group(1).strip() if business_update_match else ""

 

    return render_template('submission_output_editable.html', submission=session.get('submission', ''),

                           username=session.get('username', ''), input=input_section, output=output_section,

                           business_update=business_update_section)

 

# Define a route for updating the submission

@app.route('/update_submission', methods=['POST'])

def update_submission():

    text = session.get('submission', '')

    portfolio = session.get('portfolio', '')

    service = session.get('service', '')

    input_data = request.form.get('input')

    output_data = request.form.get('output')

    business_update = request.form.get('bu')

 

    input_pattern = r'INPUT:(.*?)OUTPUT:'

    output_pattern = r'OUTPUT:(.*?)BUSINESS UPDATE:'

    business_update_pattern = r'BUSINESS UPDATE:(.*)'

 

    # Use re.DOTALL to match across multiple lines

    input_match = re.search(input_pattern, text, re.DOTALL)

    output_match = re.search(output_pattern, text, re.DOTALL)

    business_update_match = re.search(business_update_pattern, text, re.DOTALL)

 

    # Extract the matched content

    input_section = input_match.group(1).strip() if input_match else ""

    output_section = output_match.group(1).strip() if output_match else ""

    business_update_section = business_update_match.group(1).strip() if business_update_match else ""

 

    # Convert the date to a Python datetime object

    date_column = datetime.strptime(session['selected_date'], '%Y-%m-%d')

 

    # Create a new WeekReport entry in the SQLite database

    new_report = WeekReport(

        date_column=date_column,  # Convert date to datetime

        username=session['username'],

        input_=input_data,

        output_=output_data,

        business_update=business_update,

        service=service,

        portfolio=portfolio,

        teammates=session['team'],

        progress=session['progress']

    )

 

    # Add the new report to the database

    db_session.add(new_report)

    db_session.commit()

 

    # Clear session data

    session.pop('username', None)

    session.pop('submission', None)

 

    return redirect(url_for('login'))

 

# ... (previous code)

 

# Admin Interface routes

 

@app.route('/portfolio_details', methods=['GET', 'POST'])

def portfolio_details():

    # Fetch Portfolio details

    # Fetch Portfolio details from the database

    todate = request.form['toDate']

    session['todate'] = todate

    fromdate = request.form['fromDate']

    session['fromdate'] = fromdate

    portfolio = request.form['project']

    session['portfolio'] = portfolio

    service = request.form['services']

    session['service'] = service

 

    query = ''

    df = ''

 

    if portfolio == 'all':

        query = f"SELECT * FROM weekreport WHERE DATE_COLUMN BETWEEN '{fromdate}' AND '{todate}' AND SERVICE LIKE '{service}'"

        df = pd.read_sql(query, db_session.bind)

    else:

        query = f"SELECT * FROM weekreport WHERE DATE_COLUMN BETWEEN '{fromdate}' AND '{todate}' AND PORTFOLIO LIKE '{portfolio}' AND SERVICE LIKE '{service}'"

        df = pd.read_sql(query, db_session.bind)

 

    portfolio_details = ""

 

    for num in range(0, len(df)):

        if num == 0:

            portfolio_details = portfolio_details + "\n" + f"""

            {df.iloc[num]['portfolio']}

 

            {df.iloc[num]['input_']}

            - {df.iloc[num]['output_']}

        """

        elif (num > 0 and df.iloc[num - 1]['portfolio'] != df.iloc[num]['portfolio']):

            portfolio_details = portfolio_details + "\n" + f"""

            {df.iloc[num]['portfolio']}

 

            {df.iloc[num]['input_']}

            - {df.iloc[num]['output_']}

        """

        else:

            portfolio_details = portfolio_details + "\n" + f"""

            {df.iloc[num]['input_']}

            - {df.iloc[num]['output_']}

        """

 

    return render_template_string(render_template('portfolio_details.html', portfolio_details=portfolio_details))

 

@app.route('/updated_portfolio_details', methods=['GET', 'POST'])

def update_portfolio_details():

    portfolio_details = request.form.get('portfolio-textarea')

    session["portfolio-textarea"] = portfolio_details

 

    return render_template_string(render_template('updated_portfolio_details.html', portfolio_details=portfolio_details))

 

@app.route('/download_portfolio_docx', methods=['POST'])

def download_portfolio_docx():

    portfolio_details = session.get('portfolio-textarea', '')

 

    doc = Document()

    doc.add_heading('Portfolio Details', level=1)

    doc.add_paragraph(portfolio_details)

 

    temp_docx_file = io.BytesIO()

    doc.save(temp_docx_file)

    temp_docx_file.seek(0)

 

    return send_file(

        temp_docx_file,

        as_attachment=True,

        download_name='portfolio_details.docx',

        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    )

 

@app.route('/report')

def report():

    # Define logic to generate and render the report.html page

    # This can include fetching data and rendering the template

    return render_template('report.html')

 

@app.route('/excel', methods=['POST'])

def index():

    todate = request.form['toDate']

    session['todate'] = todate

    fromdate = request.form['fromDate']

    session['fromdate'] = fromdate

    portfolio = request.form['project']

    session['portfolio'] = portfolio

    service = request.form['services']

    session['service'] = service

 

    query = ''

    df = ''

 

    if portfolio == 'all':

        query = f"SELECT * FROM weekreport WHERE DATE_COLUMN BETWEEN '{fromdate}' AND '{todate}' AND SERVICE LIKE '{service}'"

        df = pd.read_sql(query, db_session.bind)

    else:

        query = f"SELECT * FROM weekreport WHERE DATE_COLUMN BETWEEN '{fromdate}' AND '{todate}' AND PORTFOLIO LIKE '{portfolio}' AND SERVICE LIKE '{service}'"

        df = pd.read_sql(query, db_session.bind)

 

    table_html = df.to_html(classes='table table-bordered table-striped', index=False)

 

    return render_template('report.html', table_html=table_html)

 

@app.route('/download_xlsx', methods=['GET', 'POST'])

def download_xlsx():

    portfolio = session.get('portfolio', '')

    service = session.get('service', '')

    fromdate = session.get('fromdate', '')

    todate = session.get('todate', '')

 

    query = f"SELECT * FROM weekreport WHERE DATE_COLUMN BETWEEN '{fromdate}' AND '{todate}' AND PORTFOLIO LIKE '{portfolio}' AND SERVICE LIKE '{service}'"

    df = pd.read_sql(query, db_session.bind)

 

    temp_dir = tempfile.mkdtemp()

 

    xlsx_file_path = os.path.join(temp_dir, 'data.xlsx')

    with pd.ExcelWriter(xlsx_file_path, engine='xlsxwriter') as writer:

        df.to_excel(writer, sheet_name='Sheet1', index=False)

 

    return send_file(xlsx_file_path, as_attachment=True, download_name='data.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

 

@app.route('/download_pdf', methods=['POST'])

def download_pdf():

    portfolio = session.get('portfolio', '')

    service = session.get('service', '')

    fromdate = session.get('fromdate', '')

    todate = session.get('todate', '')

 

    query = f"SELECT * FROM weekreport WHERE DATE_COLUMN BETWEEN '{fromdate}' AND '{todate}' AND PORTFOLIO LIKE '{portfolio}' AND SERVICE LIKE '{service}'"

    df = pd.read_sql(query, db_session.bind)

 

    table_html = df.to_html(classes='table table-bordered table-striped', index=False)

    rendered_html = render_template('report_pdf.html', table_html=table_html)

 

    pdf_options = {

        'page-size': 'A4',

        'orientation': 'portrait',

        'no-outline': True,

        'margin-top': '0mm',

        'margin-right': '0mm',

        'margin-bottom': '0mm',

        'margin-left': '0mm',

    }

 

    pdf = pdfkit.from_string(rendered_html, False, options=pdf_options)

 

    return send_file(

        io.BytesIO(pdf),

        as_attachment=True,

        download_name='report.pdf',

        mimetype='application/pdf'

    )
    
# Run the Flask app if this file is executed directly
if __name__ == '__main__':
    app.run(debug=False , host='0.0.0.0')
